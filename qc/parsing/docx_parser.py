#!/usr/bin/env python3
"""
DOCX Parser for Qualitative Coding Interview Files

Extracts text and metadata from real interview DOCX files for processing.
Handles both AI interview transcripts and Africa interview notes formats.
"""

import os
from docx import Document
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parser for extracting text and metadata from DOCX interview files"""
    
    def __init__(self):
        """Initialize DOCX parser"""
        pass
    
    def parse_interview_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX interview file.
        
        Args:
            file_path: Path to DOCX interview file
            
        Returns:
            Dictionary containing:
            - content: Full text content
            - metadata: File metadata and properties
            - estimated_tokens: Rough token estimate
            - parsing_info: Information about parsing process
            
        Raises:
            FileNotFoundError: If file doesn't exist
            PermissionError: If can't read file
            Exception: If DOCX is corrupted or unreadable
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Interview file not found: {file_path}")
        
        if not file_path.suffix.lower() == '.docx':
            raise ValueError(f"File must be DOCX format, got: {file_path.suffix}")
        
        try:
            # Load DOCX document
            doc = Document(file_path)
            
            # Extract text content
            full_text = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # Join all text with newlines
            content = '\n'.join(full_text)
            
            # Extract document properties
            props = doc.core_properties
            
            # Calculate basic statistics
            word_count = len(content.split())
            char_count = len(content)
            line_count = len(full_text)
            
            # Create metadata
            metadata = {
                'title': props.title or file_path.stem,
                'author': props.author or 'Unknown',
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'comments': props.comments or '',
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size_bytes': file_path.stat().st_size,
                'word_count': word_count,
                'character_count': char_count,
                'line_count': line_count,
                'paragraph_count': paragraph_count,
                'parsed_timestamp': datetime.now().isoformat()
            }
            
            # Estimate tokens (will be refined with tiktoken)
            estimated_tokens = self._estimate_tokens(content)
            
            # Parsing information
            parsing_info = {
                'success': True,
                'content_length': len(content),
                'empty_paragraphs_skipped': len(doc.paragraphs) - paragraph_count,
                'extraction_method': 'paragraph_based',
                'potential_speaker_lines': self._count_potential_speakers(full_text)
            }
            
            logger.info(f"Successfully parsed {file_path.name}: "
                       f"{word_count} words, {estimated_tokens} estimated tokens")
            
            return {
                'content': content,
                'metadata': metadata,
                'estimated_tokens': estimated_tokens,
                'parsing_info': parsing_info
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {str(e)}")
            
            # Return error information
            return {
                'content': '',
                'metadata': {
                    'file_path': str(file_path),
                    'file_name': file_path.name,
                    'file_size_bytes': file_path.stat().st_size if file_path.exists() else 0,
                    'word_count': 0,
                    'character_count': 0,
                    'parsed_timestamp': datetime.now().isoformat()
                },
                'estimated_tokens': 0,
                'parsing_info': {
                    'success': False,
                    'error': str(e),
                    'error_type': type(e).__name__
                }
            }
    
    def _estimate_tokens(self, text: str) -> int:
        """
        Rough token estimation (will be replaced with tiktoken).
        Uses conservative estimate of 1.3 tokens per word.
        """
        word_count = len(text.split())
        # Conservative estimate: 1.3 tokens per word
        return int(word_count * 1.3)
    
    def _count_potential_speakers(self, paragraphs: List[str]) -> int:
        """
        Count lines that might indicate speakers.
        Looks for common patterns like "Name:", "[Speaker]", etc.
        """
        speaker_count = 0
        speaker_patterns = [
            ':',  # "John:", "Interviewer:"
            '[',  # "[Participant 1]"
            'Interviewer',  # Various interviewer formats
            'Participant',  # Various participant formats
        ]
        
        for para in paragraphs:
            # Check if paragraph starts with potential speaker indicator
            for pattern in speaker_patterns:
                if pattern in para[:50]:  # Check first 50 chars
                    speaker_count += 1
                    break
        
        return speaker_count
    
    def parse_directory(self, directory_path: Path, pattern: str = "*.docx") -> List[Dict[str, Any]]:
        """
        Parse all DOCX files in a directory.
        
        Args:
            directory_path: Directory containing DOCX files
            pattern: File pattern to match (default: "*.docx")
            
        Returns:
            List of parsed interview results
        """
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        results = []
        docx_files = list(directory_path.glob(pattern))
        
        logger.info(f"Found {len(docx_files)} DOCX files in {directory_path}")
        
        for file_path in docx_files:
            try:
                result = self.parse_interview_file(file_path)
                result['batch_info'] = {
                    'directory': str(directory_path),
                    'file_index': len(results),
                    'total_files': len(docx_files)
                }
                results.append(result)
                
            except Exception as e:
                logger.error(f"Failed to parse {file_path.name}: {str(e)}")
                # Add error result to maintain file order
                results.append({
                    'content': '',
                    'metadata': {'file_path': str(file_path), 'file_name': file_path.name},
                    'estimated_tokens': 0,
                    'parsing_info': {
                        'success': False,
                        'error': str(e),
                        'error_type': type(e).__name__
                    },
                    'batch_info': {
                        'directory': str(directory_path),
                        'file_index': len(results),
                        'total_files': len(docx_files)
                    }
                })
        
        return results
    
    def get_parsing_summary(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Generate summary statistics from parsing results.
        
        Args:
            results: List of parsing results from parse_directory()
            
        Returns:
            Summary statistics
        """
        total_files = len(results)
        successful = sum(1 for r in results if r['parsing_info']['success'])
        failed = total_files - successful
        
        if successful > 0:
            successful_results = [r for r in results if r['parsing_info']['success']]
            
            total_words = sum(r['metadata']['word_count'] for r in successful_results)
            total_tokens = sum(r['estimated_tokens'] for r in successful_results)
            
            avg_words = total_words / successful
            avg_tokens = total_tokens / successful
            
            # Find largest and smallest files
            largest = max(successful_results, key=lambda r: r['metadata']['word_count'])
            smallest = min(successful_results, key=lambda r: r['metadata']['word_count'])
        else:
            total_words = total_tokens = avg_words = avg_tokens = 0
            largest = smallest = None
        
        return {
            'total_files': total_files,
            'successful_parses': successful,
            'failed_parses': failed,
            'success_rate': (successful / total_files) * 100 if total_files > 0 else 0,
            'total_words': total_words,
            'total_estimated_tokens': total_tokens,
            'average_words_per_file': avg_words,
            'average_tokens_per_file': avg_tokens,
            'largest_file': {
                'name': largest['metadata']['file_name'],
                'words': largest['metadata']['word_count'],
                'tokens': largest['estimated_tokens']
            } if largest else None,
            'smallest_file': {
                'name': smallest['metadata']['file_name'],
                'words': smallest['metadata']['word_count'],
                'tokens': smallest['estimated_tokens']
            } if smallest else None,
            'token_limit_violations': sum(
                1 for r in results 
                if r['parsing_info']['success'] and r['estimated_tokens'] > 900_000
            )
        }


# Test/demo function
def test_docx_parser():
    """Test the DOCX parser with real interview files"""
    parser = DOCXParser()
    
    # Test paths from the real data
    ai_interviews_path = Path("data/interviews/AI_Interviews_all_2025.0728/Interviews")
    africa_interviews_path = Path("data/interviews/africa_interveiws_alll_2025.0728/For Brian_cleaned notes")
    
    print("Testing DOCX Parser with Real Interview Data\n")
    
    for path_name, path in [("AI Interviews", ai_interviews_path), ("Africa Interviews", africa_interviews_path)]:
        if path.exists():
            print(f"   {path_name}: {path}")
            
            try:
                results = parser.parse_directory(path)
                summary = parser.get_parsing_summary(results)
                
                print(f"   Parsed {summary['successful_parses']}/{summary['total_files']} files")
                print(f"   Total words: {summary['total_words']:,}")
                print(f"   Total tokens: {summary['total_estimated_tokens']:,}")
                print(f"   Average tokens/file: {summary['average_tokens_per_file']:.0f}")
                
                if summary['largest_file']:
                    print(f"   Largest: {summary['largest_file']['name']} ({summary['largest_file']['tokens']:,} tokens)")
                
                if summary['smallest_file']:
                    print(f"   Smallest: {summary['smallest_file']['name']} ({summary['smallest_file']['tokens']:,} tokens)")
                
                if summary['token_limit_violations'] > 0:
                    print(f"   WARNING: {summary['token_limit_violations']} files exceed 900K token limit")
                
                print()
                
            except Exception as e:
                print(f"   ERROR parsing directory: {e}\n")
        else:
            print(f"   {path_name}: {path} (NOT FOUND)\n")


if __name__ == "__main__":
    test_docx_parser()
"""
Phase 0: Schema Parser for User-Provided Definitions
Converts informal text/docx files into structured schemas for closed/mixed approaches
"""

import os
from pathlib import Path
from typing import Optional, Union
import docx
from pydantic import BaseModel

from .code_first_schemas import (
    ParsedCodeSchema,
    ParsedCodeDefinition,
    ParsedSpeakerSchema,
    ParsedSpeakerProperty,
    ParsedEntitySchema,
    ParsedEntityType,
    ParsedRelationshipType
)
from ..llm.llm_handler import LLMHandler


class SchemaParser:
    """Parse user-provided schema files into structured formats using LLM"""
    
    def __init__(self, llm_handler: Optional[LLMHandler] = None):
        """Initialize with LLM handler"""
        self.llm = llm_handler or LLMHandler()
    
    async def parse_code_schema(self, file_path: str) -> ParsedCodeSchema:
        """
        Parse user-provided code definitions from text/docx file.
        Handles hierarchical codes and informal descriptions.
        """
        # Read file content
        content = self._read_file(file_path)
        
        # Build prompt for LLM to parse codes
        prompt = f"""
        You are analyzing a document that contains thematic code definitions for qualitative analysis.
        The codes may be hierarchical (parent-child relationships) or flat.
        Extract all codes with their names, descriptions, and hierarchy.
        
        Document content:
        {content}
        
        Extract the following information:
        1. Code names and descriptions
        2. Hierarchical relationships (if any parent-child structure exists)
        3. Example quotes or contexts mentioned for each code (if provided)
        
        For hierarchical codes:
        - Identify parent codes (top-level themes)
        - Identify child codes (sub-themes)
        - Preserve the hierarchy structure
        
        Even if the document uses informal language like "codes related to X" or 
        "themes about Y", extract them as formal code definitions.
        
        Return a structured list of codes with:
        - name: The code name
        - description: What this code represents
        - parent_name: Name of parent code if this is a child code (null for top-level)
        - examples: Any example quotes or contexts mentioned
        """
        
        # Extract structured codes
        result = await self.llm.extract_structured(
            prompt=prompt,
            schema=ParsedCodeSchema,
            max_tokens=None  # Use maximum available
        )
        
        # Calculate hierarchy depth
        result.hierarchy_depth = self._calculate_hierarchy_depth(result.codes)
        
        return result
    
    async def parse_speaker_schema(self, file_path: str) -> ParsedSpeakerSchema:
        """
        Parse user-provided speaker property definitions.
        Identifies what properties to track for speakers.
        """
        content = self._read_file(file_path)
        
        prompt = f"""
        You are analyzing a document that describes speaker properties to track in interviews.
        Extract the schema for speaker information - NOT the actual speaker values.
        
        Document content:
        {content}
        
        Identify properties that should be tracked for each speaker, such as:
        - Demographic properties (role, organization, seniority, etc.)
        - Categorical properties (department, expertise area, etc.)
        - Descriptive properties (background, experience level, etc.)
        
        For each property determine:
        1. Property name (e.g., "role", "organization")
        2. Property type:
           - "string" for text values
           - "list" for multiple values
           - "number" for numeric values
           - "categorical" for limited set of options
        3. Whether it's required or optional
        4. Allowed values (if it's categorical)
        5. Description of what this property captures
        
        Examples of what to extract:
        - If document says "track each speaker's job title" → property: "job_title", type: "string"
        - If document says "note if speaker is junior/mid/senior" → property: "seniority", type: "categorical", allowed_values: ["junior", "mid", "senior"]
        
        Return the schema structure, not actual speaker data.
        """
        
        result = await self.llm.extract_structured(
            prompt=prompt,
            schema=ParsedSpeakerSchema,
            max_tokens=None
        )
        
        return result
    
    async def parse_entity_schema(self, file_path: str) -> ParsedEntitySchema:
        """
        Parse user-provided entity and relationship type definitions.
        Identifies types of entities and relationships to extract.
        """
        content = self._read_file(file_path)
        
        prompt = f"""
        You are analyzing a document that describes entity types and relationship types for knowledge extraction.
        Extract the schema for entities and relationships - NOT actual entities/relationships.
        
        Document content:
        {content}
        
        Extract two categories:
        
        1. ENTITY TYPES - Categories of entities to identify:
           - Entity type name (e.g., "Person", "Organization", "Technology", "Method")
           - Description of what qualifies as this entity type
           - Patterns or keywords that suggest this entity type
           
        2. RELATIONSHIP TYPES - Types of connections between entities:
           - Relationship name (e.g., "USES", "DEVELOPS", "WORKS_FOR", "COLLABORATES_WITH")
           - Description of what this relationship means
           - Valid source entity types (what types can be the source)
           - Valid target entity types (what types can be the target)
           - Whether it's directional (A→B) or bidirectional (A↔B)
        
        Examples:
        - If document mentions "track all people mentioned" → Entity type: "Person"
        - If document mentions "identify who uses which tools" → Relationship: "USES" (Person → Tool)
        - If document mentions "partnerships between organizations" → Relationship: "PARTNERS_WITH" (bidirectional)
        
        Focus on the TYPES/CATEGORIES, not specific instances.
        Default to directional relationships unless explicitly bidirectional.
        """
        
        result = await self.llm.extract_structured(
            prompt=prompt,
            schema=ParsedEntitySchema,
            max_tokens=None
        )
        
        return result
    
    def _read_file(self, file_path: str) -> str:
        """Read content from text or docx file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Handle different file types
        if path.suffix.lower() == '.docx':
            return self._read_docx(file_path)
        elif path.suffix.lower() in ['.txt', '.text']:
            return self._read_text(file_path)
        else:
            # Try to read as text by default
            return self._read_text(file_path)
    
    def _read_docx(self, file_path: str) -> str:
        """Read content from DOCX file"""
        doc = docx.Document(file_path)
        
        # Extract all paragraphs
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n\n".join(content)
    
    def _read_text(self, file_path: str) -> str:
        """Read content from text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _calculate_hierarchy_depth(self, codes: list[ParsedCodeDefinition]) -> int:
        """Calculate the maximum depth of code hierarchy"""
        if not codes:
            return 0
        
        # Build parent-child mapping
        children_map = {}
        root_codes = []
        
        for code in codes:
            if code.parent_name:
                if code.parent_name not in children_map:
                    children_map[code.parent_name] = []
                children_map[code.parent_name].append(code.name)
            else:
                root_codes.append(code.name)
        
        # Calculate max depth
        def get_depth(code_name: str, current_depth: int = 1) -> int:
            if code_name not in children_map:
                return current_depth
            
            max_child_depth = current_depth
            for child in children_map[code_name]:
                child_depth = get_depth(child, current_depth + 1)
                max_child_depth = max(max_child_depth, child_depth)
            
            return max_child_depth
        
        max_depth = 0
        for root in root_codes:
            depth = get_depth(root)
            max_depth = max(max_depth, depth)
        
        return max_depth if max_depth > 0 else 1


class SchemaValidator:
    """Validate parsed schemas for completeness and consistency"""
    
    @staticmethod
    def validate_code_schema(schema: ParsedCodeSchema) -> tuple[bool, list[str]]:
        """Validate code schema for issues"""
        issues = []
        
        # Check for empty codes
        if not schema.codes:
            issues.append("No codes found in schema")
            return False, issues
        
        # Check for duplicate code names
        code_names = [code.name for code in schema.codes]
        if len(code_names) != len(set(code_names)):
            duplicates = [name for name in code_names if code_names.count(name) > 1]
            issues.append(f"Duplicate code names found: {duplicates}")
        
        # Check for orphaned parent references
        all_names = set(code_names)
        for code in schema.codes:
            if code.parent_name and code.parent_name not in all_names:
                issues.append(f"Code '{code.name}' references non-existent parent '{code.parent_name}'")
        
        # Check for circular dependencies
        parent_map = {code.name: code.parent_name for code in schema.codes if code.parent_name}
        for code_name in parent_map:
            visited = set()
            current = code_name
            while current:
                if current in visited:
                    issues.append(f"Circular dependency detected involving code '{code_name}'")
                    break
                visited.add(current)
                current = parent_map.get(current)
        
        return len(issues) == 0, issues
    
    @staticmethod
    def validate_speaker_schema(schema: ParsedSpeakerSchema) -> tuple[bool, list[str]]:
        """Validate speaker schema for issues"""
        issues = []
        
        # Check for empty properties
        if not schema.properties:
            issues.append("No speaker properties found in schema")
            return False, issues
        
        # Check for duplicate property names
        prop_names = [prop.name for prop in schema.properties]
        if len(prop_names) != len(set(prop_names)):
            duplicates = [name for name in prop_names if prop_names.count(name) > 1]
            issues.append(f"Duplicate property names found: {duplicates}")
        
        # Validate property types
        valid_types = ["string", "list", "number", "categorical"]
        for prop in schema.properties:
            if prop.property_type not in valid_types:
                issues.append(f"Invalid property type '{prop.property_type}' for property '{prop.name}'")
            
            # Check categorical properties have allowed values
            if prop.property_type == "categorical" and not prop.allowed_values:
                issues.append(f"Categorical property '{prop.name}' missing allowed values")
        
        return len(issues) == 0, issues
    
    @staticmethod
    def validate_entity_schema(schema: ParsedEntitySchema) -> tuple[bool, list[str]]:
        """Validate entity/relationship schema for issues"""
        issues = []
        
        # Check for empty types
        if not schema.entity_types:
            issues.append("No entity types found in schema")
        
        if not schema.relationship_types:
            issues.append("Warning: No relationship types found in schema")
        
        # Check for duplicate entity type names
        entity_names = [et.name for et in schema.entity_types]
        if len(entity_names) != len(set(entity_names)):
            duplicates = [name for name in entity_names if entity_names.count(name) > 1]
            issues.append(f"Duplicate entity type names found: {duplicates}")
        
        # Check for duplicate relationship type names
        rel_names = [rt.name for rt in schema.relationship_types]
        if len(rel_names) != len(set(rel_names)):
            duplicates = [name for name in rel_names if rel_names.count(name) > 1]
            issues.append(f"Duplicate relationship type names found: {duplicates}")
        
        # Validate relationship source/target types exist
        valid_entity_types = set(entity_names)
        for rel in schema.relationship_types:
            for source_type in rel.valid_source_types:
                if source_type not in valid_entity_types:
                    issues.append(f"Relationship '{rel.name}' references non-existent source type '{source_type}'")
            
            for target_type in rel.valid_target_types:
                if target_type not in valid_entity_types:
                    issues.append(f"Relationship '{rel.name}' references non-existent target type '{target_type}'")
        
        return len(issues) == 0, issues